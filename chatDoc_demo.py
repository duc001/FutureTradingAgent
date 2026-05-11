import os

os.environ['HF_HUB_DISABLE_SYMLINKS_WARNING'] = '1'
os.environ['TRANSFORMERS_OFFLINE'] = '0'

from langchain_chroma import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangChainDocument

from docx import Document as WordDocument
from PyPDF2 import PdfReader
from langchain_community.document_loaders import UnstructuredExcelLoader

DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "sk-205ee372c5df491f8050324eb697e504")
DEEPSEEK_BASE_URL = "https://api.deepseek.com"
DEEPSEEK_MODEL_NAME = "deepseek-v4-flash"
TEMPERATURE = 0.7
FILE_NAME = "2026-test.docx"
PDF_FILE = "2026-test.pdf"
EXCEL_FILE = "跟踪止损策略回测数据_202605.xlsx"
CHROMA_DB_PATH = "./chroma_db"
COLLECTION_NAME = "chatdoc_collection"
EMBEDDING_MODEL_NAME = "paraphrase-multilingual-MiniLM-L12-v2"


class ChatDoc:
    """统一文档读取器 - 支持 Word、PDF、Excel"""

    READERS = {
        '.docx': '_read_word',
        '.pdf': '_read_pdf',
        '.xlsx': '_read_excel',
        '.xls': '_read_excel'
    }

    def __init__(self, file_name=None):
        """初始化

        Args:
            file_name: 文档路径
        """
        self.file_name = file_name if file_name else FILE_NAME

        if not os.path.exists(self.file_name):
            raise FileNotFoundError(f"文件不存在: {self.file_name}")

        ext = os.path.splitext(self.file_name)[1].lower()
        if ext not in self.READERS:
            raise ValueError(f"不支持的文件格式: {ext}\n支持的格式: {list(self.READERS.keys())}")

        self.reader_method = self.READERS[ext]

    def get_file_content(self):
        """获取文档内容（用于 RAG 处理）

        Returns:
            List[LangChainDocument]: LangChain 文档对象列表，包含内容和元数据
        """
        reader_func = getattr(self, self.reader_method)
        return reader_func()

    def _read_word(self):
        """读取 Word 文档"""
        doc = WordDocument(self.file_name)
        text = "\n".join([para.text for para in doc.paragraphs])
        return [LangChainDocument(
            page_content=text,
            metadata={"source": self.file_name, "type": "word"}
        )]

    def _read_pdf(self):
        """读取 PDF 文档"""
        reader = PdfReader(self.file_name)
        documents = []

        for page_num, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text()
            if page_text:
                doc = LangChainDocument(
                    page_content=page_text,
                    metadata={"source": self.file_name, "type": "pdf", "page": page_num}
                )
                documents.append(doc)

        return documents

    def _read_excel(self):
        """读取 Excel 文档"""
        loader = UnstructuredExcelLoader(self.file_name)
        documents = loader.load()

        for doc in documents:
            doc.metadata["source"] = self.file_name
            doc.metadata["type"] = "excel"

        return documents

    def split_data(self, chunk_size=800, chunk_overlap=100):
        """切分数据，用于 RAG 系统

        Args:
            chunk_size: 每个块的大小（字符数）
            chunk_overlap: 块之间的重叠大小
        """
        documents = self.get_file_content()
        if documents is not None and len(documents) > 0:
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", " ", ""]
            )
            split_docs = splitter.split_documents(documents)
            print(f"切分: {len(documents)} -> {len(split_docs)} 块")

            return split_docs
        else:
            print("警告: 未获取到文档内容")
            return []

    def embed_and_store(self, persist_directory=CHROMA_DB_PATH):
        """将文档块向量化并存储到 Chroma 数据库

        Args:
            persist_directory: Chroma 数据库持久化路径

        Returns:
            Chroma: LangChain Chroma 向量存储对象
        """
        print("\n" + "=" * 50)
        print("[阶段 1] 向量化与存储")
        print("=" * 50)

        split_docs = self.split_data()

        if not split_docs:
            print("错误: 没有可处理的文档")
            return None

        print("\n加载嵌入模型...")
        try:
            from langchain_huggingface import HuggingFaceEmbeddings
            embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL_NAME)
        except Exception as e:
            print(f"模型加载失败: {e}")
            return None

        print("正在构建向量数据库...")
        db = Chroma.from_documents(
            documents=split_docs,
            embedding=embeddings,
            persist_directory=persist_directory,
            collection_name=COLLECTION_NAME
        )
        print(f"存储完成: 共 {len(split_docs)} 个向量块")
        print(f"   路径: {persist_directory}")
        print(f"   集合: {COLLECTION_NAME}")

        return db

    def ask_answer(self, query="唐奇安策略是什么？", top_k=6):
        """标准相似度搜索
        
        Args:
            query: 查询问题
            top_k: 返回结果数量
        """
        print(f"\n提问: {query}")
        db = self.embed_and_store()
        if not db: return []

        results = db.similarity_search(query, k=top_k)
        
        if results:
            print(f"检索到 {len(results)} 条结果")
            for i, doc in enumerate(results[:3], 1):
                clean_content = doc.page_content.replace('\n', ' ').replace('\t', ' ')
                preview = clean_content[:80] + '...' if len(clean_content) > 80 else clean_content
                print(f"   [{i}] {preview}")
        
        return results

    def ask_answer_with_score(self, query="唐奇安策略是什么？", top_k=6):
        """带分数的相似度搜索
        
        Args:
            query: 查询问题
            top_k: 返回结果数量
        """
        print(f"\n提问: {query}")
        db = self.embed_and_store()
        if not db: return []

        results = db.similarity_search_with_score(query, k=top_k)
        
        if results:
            print(f"检索到 {len(results)} 条结果")
            for i, (doc, score) in enumerate(results[:3], 1):
                clean_content = doc.page_content.replace('\n', ' ').replace('\t', ' ')
                preview = clean_content[:80] + '...' if len(clean_content) > 80 else clean_content
                print(f"   [{i}] 分数: {score:.4f}")
                print(f"       内容: {preview}")
        
        res = [doc for doc, score in results]
        return res

    def ask_answer_with_mmr(self, query="唐奇安策略是什么？", top_k=6, fetch_k=20):
        """使用 MMR（最大边际相关性）检索

        Args:
            query: 查询问题
            top_k: 返回结果数量
            fetch_k: 初始候选集大小
        """
        print(f"\n提问: {query}")
        db = self.embed_and_store()
        if not db: return []

        retriever = db.as_retriever(
            search_type="mmr",
            search_kwargs={"k": top_k, "fetch_k": fetch_k}
        )

        results = retriever.invoke(query)

        if results:
            print(f"MMR 检索到 {len(results)} 条结果 (fetch_k={fetch_k})")
            for i, doc in enumerate(results[:3], 1):
                clean_content = doc.page_content.replace('\n', ' ').replace('\t', ' ')
                preview = clean_content[:80] + '...' if len(clean_content) > 80 else clean_content
                print(f"   [{i}] {preview}")

        return results

    def _call_llm(self, system_prompt, user_prompt, temperature=None, max_tokens=800, task_name="LLM"):
        """统一的 LLM 调用方法
        
        Args:
            system_prompt: 系统提示词
            user_prompt: 用户提示词
            temperature: 温度参数
            max_tokens: 最大 token 数
            task_name: 任务名称（用于日志）
        
        Returns:
            str: LLM 生成的回答
        """
        try:
            from langchain_openai import ChatOpenAI
            from langchain_core.messages import SystemMessage, HumanMessage
            
            temp = temperature if temperature is not None else TEMPERATURE
            
            llm = ChatOpenAI(
                model=DEEPSEEK_MODEL_NAME,
                api_key=DEEPSEEK_API_KEY,
                base_url=DEEPSEEK_BASE_URL,
                temperature=temp,
                max_tokens=max_tokens,
            )
            
            messages = [
                SystemMessage(content=system_prompt),
                HumanMessage(content=user_prompt)
            ]
            
            print(f"\n正在调用 {task_name}...")
            response = llm.invoke(messages)
            
            print(f"\n{'='*70}")
            print(f"{task_name} 结果")
            print(f"{'='*70}")
            print(response.content)
            print(f"{'='*70}")
            
            return response.content
            
        except Exception as e:
            print(f"{task_name} 调用失败: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def chat_with_llm(self, query="唐奇安策略是什么？", top_k=6, search_type="standard", temperature=None):
        """使用 LLM 分析检索结果并生成回答
        
        Args:
            query: 用户问题
            top_k: 检索文档数量
            search_type: 检索方式 ('standard', 'score', 'mmr')
            temperature: 温度参数
        
        Returns:
            str: LLM 生成的回答
        """
        print(f"\n提问: {query}")
        print(f"检索方式: {search_type}")
        
        analysis_keywords = ["分类", "汇总", "总结", "概括", "归纳", "整理", "梳理"]
        statistics_keywords = ["统计", "数据", "数量", "频率", "分布"]
        
        is_analysis_task = any(keyword in query for keyword in analysis_keywords)
        is_statistics_task = any(keyword in query for keyword in statistics_keywords)
        
        if is_statistics_task:
            print("检测到统计分析类任务，使用全文分析模式")
            return self.analyze_full_document(query, temperature, is_statistics=True)
        elif is_analysis_task:
            print("检测到全文分析类任务，使用全文分析模式")
            return self.analyze_full_document(query, temperature, is_statistics=False)
        
        if search_type == "standard":
            docs = self.ask_answer(query, top_k=top_k)
        elif search_type == "score":
            docs = self.ask_answer_with_score(query, top_k=top_k)
        elif search_type == "mmr":
            docs = self.ask_answer_with_mmr(query, top_k=top_k, fetch_k=20)
        else:
            raise ValueError(f"不支持的检索方式: {search_type}")
        
        if not docs:
            return "未找到相关文档"
        
        context_parts = []
        for i, doc in enumerate(docs, 1):
            source_info = doc.metadata.get('source', '未知来源')
            page_info = doc.metadata.get('page', '')
            location = f"{source_info}"
            if page_info:
                location += f" (第{page_info}页)"
            
            context_parts.append(
                f"【文档{i}】来源: {location}\n{doc.page_content.strip()}"
            )
        
        context = "\n\n".join(context_parts)
        
        system_prompt = """你是一个专业的智能文档助手，擅长基于检索到的文档准确回答问题。

## 回答原则
1. **严格基于文档**: 只使用提供的文档内容作答，不要编造信息
2. **引用来源**: 回答时标注信息来源，如“根据文档1..."
3. **诚实承认**: 如果文档中没有相关信息，明确说明“文档中未找到相关信息”
4. **结构化回答**: 使用清晰的条理和格式组织答案
5. **简洁准确**: 避免冗余，直击要点

## 回答格式
- 先给出核心结论
- 再展开详细说明
- 引用相关文档支撑观点
- 如有多个角度，分点阐述"""
        
        user_prompt = f"""请基于以下相关文档回答问题。

## 用户问题
{query}

## 相关文档
{context}

## 你的回答
请基于以上文档，给出准确、结构化的回答："""
        
        return self._call_llm(system_prompt, user_prompt, temperature, max_tokens=800, task_name="RAG 问答")
    
    def analyze_full_document(self, query, temperature=None, is_statistics=False):
        """全文分析模式 - 用于分类、汇总、统计等需要全局理解的任务
        
        Args:
            query: 分析任务描述
            temperature: 温度参数
            is_statistics: 是否为统计分析任务
        
        Returns:
            str: LLM 生成的分析结果
        """
        # 获取所有文档块
        all_docs = self.split_data(chunk_size=800, chunk_overlap=100)
        
        if not all_docs:
            return "无法获取文档内容"
        
        print(f"获取到 {len(all_docs)} 个文档块进行分析")
        
        # 限制最大块数，避免超出 token 限制
        max_chunks = 20
        if len(all_docs) > max_chunks:
            print(f"文档块较多 ({len(all_docs)})，选取前 {max_chunks} 个关键块")
            selected_docs = all_docs[:max_chunks]
        else:
            selected_docs = all_docs
        
        # 构建完整上下文
        context_parts = []
        for i, doc in enumerate(selected_docs, 1):
            page_info = doc.metadata.get('page', '')
            location = f"第{page_info}页" if page_info else f"段落{i}"
            context_parts.append(
                f"【{location}】\n{doc.page_content.strip()}"
            )
        
        full_context = "\n\n".join(context_parts)
        
        # 根据任务类型构建不同的 Prompt
        if is_statistics:
            system_prompt = """你是一个专业的数据分析专家，擅长从文档中提取统计信息和数据洞察。

## 分析原则
1. **数据驱动**: 基于文档中的实际数据进行统计分析
2. **量化表达**: 使用数字、百分比、频率等量化指标
3. **发现规律**: 识别数据中的模式、趋势和异常
4. **可视化思维**: 用表格、列表等形式呈现统计数据
5. **客观准确**: 忠实于原文数据，不夸大或缩小

## 输出格式
- 使用 Markdown 表格展示统计数据
- 列出关键指标和发现
- 提供数据解读和洞察
- 如有必要，给出建议"""
        else:
            system_prompt = """你是一个专业的文档分析专家，擅长对文档内容进行分类、汇总和结构化整理。

## 分析原则
1. **全面理解**: 基于提供的全部内容进行分析
2. **逻辑清晰**: 按照主题、类别或逻辑关系组织内容
3. **层次分明**: 使用标题、子标题、列表等结构化格式
4. **精炼概括**: 提取关键信息，去除冗余
5. **客观准确**: 忠实于原文，不添加个人观点

## 输出格式
- 使用 Markdown 格式
- 采用层级结构（一级标题、二级标题、列表）
- 每个类别下有清晰的要点
- 适当引用原文关键信息"""
        
        task_description = "统计分析" if is_statistics else "分类汇总"
        user_prompt = f"""请对以下文档内容进行{task_description}。

## 分析任务
{query}

## 文档内容
{full_context}

## 分析结果
请按照上述要求，给出结构化的{task_description}结果："""
        
        # 统一调用 _call_llm
        return self._call_llm(system_prompt, user_prompt, temperature, max_tokens=1500, task_name="全文分析")



import time

test_file = EXCEL_FILE
print("\nChatDoc RAG 系统测试")
print(f"文件: {test_file} | 模型: {EMBEDDING_MODEL_NAME}")
print("提示: 输入 'END' 结束程序\n")
start_time = time.time()

try:
    chatdoc = ChatDoc(test_file)
    print(f"加载成功: {chatdoc.file_name}")
    
    while True:
        query = input("\n请输入问题 (输入 END 结束): ").strip()
        
        if query.upper() == "END":
            print("\n感谢使用，再见！")
            break
        
        if not query:
            continue
        
        print("\n" + "="*70)
        print(f"问题: {query}")
        print("="*70)
        
        search_methods = [
            ("standard", "标准相似度搜索"),
            ("score", "带分数搜索"),
            ("mmr", "MMR 搜索")
        ]
        
        for search_type, method_name in search_methods:
            print(f"\n{'='*70}")
            print(f"【{method_name}】")
            print(f"{'='*70}")
            
            answer = chatdoc.chat_with_llm(query, top_k=6, search_type=search_type)
            
            if answer:
                print(f"\n✓ {method_name} 完成")
            else:
                print(f"\n✗ {method_name} 未能获取回答")
        
        print("\n" + "="*70)
        print("三种方式对比完成！可以继续提问或输入 END 结束。")
        print("="*70)

    print(f"\n总耗时: {time.time() - start_time:.2f}s")

except KeyboardInterrupt:
    print("\n\n程序被用户中断")
except Exception as e:
    print(f"\n异常: {e}")
    import traceback
    traceback.print_exc()